"""
Builds DOCX and PDF telemetry reports from the nested metrics-and-insights JSON
written by the Node.js report service. Keeping this renderer standalone makes
the report output reproducible and lets the API surface actionable errors when
document generation fails.
"""

from __future__ import annotations

import argparse
import json
import sys
from datetime import datetime, timezone
from html import escape
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


REQUIRED_KEYS = (
    "reportTitle",
    "deviceId",
    "executiveSummary",
    "keyMetrics",
    "insights",
    "anomalies",
    "recommendations",
)

METRIC_LABELS = (
    ("distanceKm", "Distance", " km"),
    ("engineRunMinutes", "Engine run time", " min"),
    ("idleMinutes", "Idle time", " min"),
    ("stops", "Stops", ""),
    ("avgSpeed", "Average speed", " km/h"),
    ("maxSpeed", "Maximum speed", " km/h"),
)

DETAIL_COLUMNS = (
    ("Start", "windowStart"),
    ("End", "windowEnd"),
    ("Distance km", "distanceKm"),
    ("Engine min", "engineRunMinutes"),
    ("Idle min", "idleMinutes"),
    ("Stops", "stops"),
    ("Avg speed", "avgSpeed"),
    ("Max speed", "maxSpeed"),
)

SEVERITY_COLORS = {
    "low": (RGBColor(52, 89, 64), colors.HexColor("#345940")),
    "medium": (RGBColor(165, 95, 0), colors.HexColor("#A55F00")),
    "high": (RGBColor(170, 35, 35), colors.HexColor("#AA2323")),
    "critical": (RGBColor(130, 15, 15), colors.HexColor("#820F0F")),
}


def normalize_data(raw: dict[str, Any]) -> dict[str, Any]:
    """Converts Phase 3's {metrics, insights} object into renderer fields."""
    metrics = raw.get("metrics", raw)
    insights = raw.get("insights", raw)
    report_period = metrics.get("reportPeriod", raw.get("reportPeriod", {}))

    return {
        "deviceId": metrics.get("deviceId", raw.get("deviceId")),
        "reportTitle": insights.get("reportTitle", raw.get("reportTitle")),
        "periodStart": report_period.get("from", raw.get("periodStart")),
        "periodEnd": report_period.get("to", raw.get("periodEnd")),
        "executiveSummary": insights.get(
            "executiveSummary", raw.get("executiveSummary")
        ),
        "keyMetrics": insights.get("keyMetrics", raw.get("keyMetrics")),
        "insights": insights.get("insights", raw.get("insights")),
        "anomalies": insights.get("anomalies", raw.get("anomalies")),
        "recommendations": insights.get(
            "recommendations", raw.get("recommendations")
        ),
        "notableWindows": insights.get(
            "notableWindows", raw.get("notableWindows", [])
        ),
        "windows": metrics.get("windows", raw.get("windows", [])),
        "totals": metrics.get("totals", raw.get("totals", {})),
    }


def validate_data(data: dict[str, Any]) -> None:
    for key in REQUIRED_KEYS:
        if key not in data or data[key] is None or data[key] == "":
            raise ValueError(f"Missing required key: {key}")

    for key in ("insights", "anomalies", "recommendations"):
        if not isinstance(data[key], list):
            raise ValueError(f"Invalid required key: {key} must be a list")

    if not isinstance(data["keyMetrics"], dict):
        raise ValueError("Invalid required key: keyMetrics must be an object")


def load_data(input_path: Path) -> dict[str, Any]:
    try:
        with input_path.open("r", encoding="utf-8") as source:
            raw = json.load(source)
    except FileNotFoundError as error:
        raise ValueError(f"Input JSON does not exist: {input_path}") from error
    except json.JSONDecodeError as error:
        raise ValueError(f"Invalid input JSON: {error.msg}") from error

    if not isinstance(raw, dict):
        raise ValueError("Invalid input JSON: expected an object")

    data = normalize_data(raw)
    validate_data(data)
    return data


def metric_value(data: dict[str, Any], key: str, suffix: str) -> str:
    value = data["keyMetrics"].get(key, data.get("totals", {}).get(key, 0))
    return f"{value}{suffix}"


def set_cell_shading(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_width(cell: Any, width_dxa: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    cell_width = properties.find(qn("w:tcW"))
    if cell_width is None:
        cell_width = OxmlElement("w:tcW")
        properties.append(cell_width)
    cell_width.set(qn("w:w"), str(width_dxa))
    cell_width.set(qn("w:type"), "dxa")


def set_table_geometry(table: Any, widths_dxa: list[int]) -> None:
    table.autofit = False
    table_properties = table._tbl.tblPr
    table_width = table_properties.first_child_found_in("w:tblW")
    table_width.set(qn("w:w"), str(sum(widths_dxa)))
    table_width.set(qn("w:type"), "dxa")

    table_indent = table_properties.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_properties.append(table_indent)
    table_indent.set(qn("w:w"), "120")
    table_indent.set(qn("w:type"), "dxa")

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)


def format_docx_table_header(row: Any) -> None:
    for cell in row.cells:
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(31, 77, 120)


def add_docx_heading(document: Document, text: str) -> None:
    document.add_heading(text, level=1)


def add_docx_bullets(document: Document, items: list[Any], item_renderer) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        item_renderer(paragraph, item)


def configure_docx(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading = document.styles["Heading 1"]
    heading.font.name = "Calibri"
    heading.font.size = Pt(16)
    heading.font.color.rgb = RGBColor(46, 116, 181)
    heading.paragraph_format.space_before = Pt(16)
    heading.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("Telemetry Analytics Report")
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = RGBColor(99, 115, 129)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(
        f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(99, 115, 129)


def build_docx(data: dict[str, Any], output_path: Path) -> None:
    document = Document()
    configure_docx(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run(data["reportTitle"])
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    subtitle_run = subtitle.add_run(
        f"Device {data['deviceId']} | {data.get('periodStart', 'N/A')} to "
        f"{data.get('periodEnd', 'N/A')}"
    )
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(85, 85, 85)

    add_docx_heading(document, "Executive Summary")
    document.add_paragraph(data["executiveSummary"])

    add_docx_heading(document, "Key Metrics")
    metric_table = document.add_table(rows=1, cols=2, style="Table Grid")
    metric_table.rows[0].cells[0].text = "Metric"
    metric_table.rows[0].cells[1].text = "Value"
    format_docx_table_header(metric_table.rows[0])
    for key, label, suffix in METRIC_LABELS:
        cells = metric_table.add_row().cells
        cells[0].text = label
        cells[1].text = metric_value(data, key, suffix)
    set_table_geometry(metric_table, [2700, 6660])

    add_docx_heading(document, "Insights")
    def render_insight(paragraph: Any, insight: Any) -> None:
        if isinstance(insight, dict):
            title_run = paragraph.add_run(f"{insight.get('title', 'Insight')}: ")
            title_run.bold = True
            paragraph.add_run(str(insight.get("detail", "")))
        else:
            paragraph.add_run(str(insight))
    add_docx_bullets(document, data["insights"], render_insight)

    add_docx_heading(document, "Anomalies")
    anomaly_table = document.add_table(rows=1, cols=2, style="Table Grid")
    anomaly_table.rows[0].cells[0].text = "Severity"
    anomaly_table.rows[0].cells[1].text = "Observation"
    format_docx_table_header(anomaly_table.rows[0])
    for anomaly in data["anomalies"]:
        anomaly = anomaly if isinstance(anomaly, dict) else {"detail": str(anomaly)}
        severity = str(anomaly.get("severity", "low")).lower()
        cells = anomaly_table.add_row().cells
        cells[0].text = severity.title()
        cells[1].text = (
            f"{anomaly.get('title', 'Observation')}: {anomaly.get('detail', '')}"
        )
        color, _ = SEVERITY_COLORS.get(severity, SEVERITY_COLORS["low"])
        for cell in cells:
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = color
                run.bold = severity in {"medium", "high", "critical"}
    if not data["anomalies"]:
        cells = anomaly_table.add_row().cells
        cells[0].text = "None"
        cells[1].text = "No anomalies were identified for this report period."
    set_table_geometry(anomaly_table, [1800, 7560])

    add_docx_heading(document, "Recommendations")
    add_docx_bullets(document, data["recommendations"], lambda p, item: p.add_run(str(item)))

    notable_windows = data.get("notableWindows", [])
    if notable_windows:
        add_docx_heading(document, "Notable Windows")
        def render_notable(paragraph: Any, window: Any) -> None:
            window = window if isinstance(window, dict) else {"summary": str(window)}
            period_run = paragraph.add_run(
                f"{window.get('windowStart', 'N/A')} to {window.get('windowEnd', 'N/A')}: "
            )
            period_run.bold = True
            paragraph.add_run(str(window.get("summary", "")))
        add_docx_bullets(document, notable_windows, render_notable)

    windows = data.get("windows", [])
    if windows:
        add_docx_heading(document, "Window-by-Window Detail")
        detail_table = document.add_table(rows=1, cols=len(DETAIL_COLUMNS), style="Table Grid")
        for index, (label, _) in enumerate(DETAIL_COLUMNS):
            detail_table.rows[0].cells[index].text = label
        format_docx_table_header(detail_table.rows[0])
        for window in windows:
            cells = detail_table.add_row().cells
            for index, (_, key) in enumerate(DETAIL_COLUMNS):
                cells[index].text = str(window.get(key, ""))
        set_table_geometry(detail_table, [1350, 1350, 1050, 1125, 975, 675, 1395, 1440])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def pdf_paragraph_style(base: ParagraphStyle, **overrides: Any) -> ParagraphStyle:
    values = {"parent": base}
    values.update(overrides)
    return ParagraphStyle(f"custom_{len(overrides)}_{id(overrides)}", **values)


def build_pdf(data: dict[str, Any], output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ReportTitle", parent=styles["Title"], fontName="Helvetica-Bold",
        fontSize=22, leading=26, textColor=colors.HexColor("#0B2545"), alignment=TA_CENTER,
        spaceAfter=4,
    )
    subtitle_style = ParagraphStyle(
        "ReportSubtitle", parent=styles["Normal"], fontSize=9, leading=12,
        textColor=colors.HexColor("#555555"), alignment=TA_CENTER, spaceAfter=16,
    )
    heading_style = ParagraphStyle(
        "ReportHeading", parent=styles["Heading1"], fontName="Helvetica-Bold",
        fontSize=15, leading=19, textColor=colors.HexColor("#2E74B5"),
        spaceBefore=12, spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "ReportBody", parent=styles["BodyText"], fontSize=10, leading=13, spaceAfter=5,
    )
    story: list[Any] = [
        Paragraph(escape(data["reportTitle"]), title_style),
        Paragraph(
            escape(
                f"Device {data['deviceId']} | {data.get('periodStart', 'N/A')} to "
                f"{data.get('periodEnd', 'N/A')}"
            ),
            subtitle_style,
        ),
        Paragraph("Executive Summary", heading_style),
        Paragraph(escape(data["executiveSummary"]), body_style),
        Paragraph("Key Metrics", heading_style),
    ]

    metric_rows = [["Metric", "Value"]] + [
        [label, metric_value(data, key, suffix)]
        for key, label, suffix in METRIC_LABELS
    ]
    metric_table = Table(metric_rows, colWidths=[2.0 * inch, 4.5 * inch])
    metric_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1F4D78")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#B8C4D0")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.extend([metric_table, Paragraph("Insights", heading_style)])

    for insight in data["insights"]:
        if isinstance(insight, dict):
            text = f"<b>{escape(str(insight.get('title', 'Insight')))}:</b> {escape(str(insight.get('detail', '')))}"
        else:
            text = escape(str(insight))
        story.append(Paragraph(f"- {text}", body_style))

    story.append(Paragraph("Anomalies", heading_style))
    anomaly_rows = [["Severity", "Observation"]]
    anomaly_colors: list[Any] = []
    for anomaly in data["anomalies"]:
        anomaly = anomaly if isinstance(anomaly, dict) else {"detail": str(anomaly)}
        severity = str(anomaly.get("severity", "low")).lower()
        anomaly_rows.append([
            severity.title(),
            f"{anomaly.get('title', 'Observation')}: {anomaly.get('detail', '')}",
        ])
        anomaly_colors.append(SEVERITY_COLORS.get(severity, SEVERITY_COLORS["low"])[1])
    if len(anomaly_rows) == 1:
        anomaly_rows.append(["None", "No anomalies were identified for this report period."])
        anomaly_colors.append(colors.HexColor("#345940"))
    anomaly_table = Table(anomaly_rows, colWidths=[1.25 * inch, 5.25 * inch])
    anomaly_style = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1F4D78")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#B8C4D0")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]
    for index, color in enumerate(anomaly_colors, start=1):
        anomaly_style.extend([
            ("TEXTCOLOR", (0, index), (-1, index), color),
            ("FONTNAME", (0, index), (-1, index), "Helvetica-Bold" if index else "Helvetica"),
        ])
    anomaly_table.setStyle(TableStyle(anomaly_style))
    story.extend([anomaly_table, Paragraph("Recommendations", heading_style)])
    for recommendation in data["recommendations"]:
        story.append(Paragraph(f"- {escape(str(recommendation))}", body_style))

    if data.get("notableWindows"):
        story.append(Paragraph("Notable Windows", heading_style))
        for window in data["notableWindows"]:
            window = window if isinstance(window, dict) else {"summary": str(window)}
            story.append(Paragraph(
                escape(
                    f"- {window.get('windowStart', 'N/A')} to {window.get('windowEnd', 'N/A')}: "
                    f"{window.get('summary', '')}"
                ),
                body_style,
            ))

    if data.get("windows"):
        story.append(Paragraph("Window-by-Window Detail", heading_style))
        detail_rows = [[label for label, _ in DETAIL_COLUMNS]]
        for window in data["windows"]:
            detail_rows.append([str(window.get(key, "")) for _, key in DETAIL_COLUMNS])
        detail_table = Table(
            detail_rows,
            colWidths=[0.82 * inch, 0.82 * inch, 0.6 * inch, 0.68 * inch,
                       0.58 * inch, 0.42 * inch, 0.82 * inch, 0.76 * inch],
            repeatRows=1,
        )
        detail_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1F4D78")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#B8C4D0")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 3),
            ("RIGHTPADDING", (0, 0), (-1, -1), 3),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(detail_table)

    story.extend([
        Spacer(1, 10),
        Paragraph(
            escape(f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"),
            pdf_paragraph_style(body_style, fontSize=8, textColor=colors.HexColor("#637381")),
        ),
    ])
    document.build(story)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Generate a telemetry report document")
    parser.add_argument("--input", required=True, help="Combined report JSON input path")
    parser.add_argument("--output", required=True, help="Generated document output path")
    parser.add_argument("--format", required=True, choices=("docx", "pdf"))
    args = parser.parse_args(argv)

    try:
        data = load_data(Path(args.input))
        output_path = Path(args.output)
        if args.format == "docx":
            build_docx(data, output_path)
        else:
            build_pdf(data, output_path)

        print(f"Generated {args.format.upper()} report: {output_path} ({output_path.stat().st_size} bytes)")
        return 0
    except Exception as error:
        print(f"Report generation failed: {error}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
